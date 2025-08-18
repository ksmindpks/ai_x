# -*- coding: utf-8 -*-
"""
reindex_upstage_docx.py
- input_dirs 내의 모든 .docx 파일을 재임베딩하여 Pinecone 인덱스(codedoc-law-upstage)에 업서트
- 임베딩: Upstage solar-embedding-1-large (4096차원)
- 벡터DB: Pinecone (cosine)
- 배치/리트라이/진행로그 포함

필수 패키지:
    pip install python-docx langchain-upstage pinecone-client python-dotenv

필수 환경변수:
    UPSTAGE_API_KEY
    PINECONE_API_KEY
선택 환경변수:
    PINECONE_CLOUD=aws
    PINECONE_REGION=us-east-1
    NAMESPACE=optional-namespace
"""

from __future__ import annotations
import os
import sys
import time
import glob
import math
import json
import hashlib
from typing import Iterable, Iterator, List, Tuple, Dict

try:
    from dotenv import load_dotenv  # 선택
    load_dotenv()
except Exception:
    pass

from pinecone import Pinecone, ServerlessSpec
from langchain_upstage import UpstageEmbeddings
from docx import Document

# ===================== 사용자 설정 =====================

# 대상 디렉토리들 (질문에 주신 경로)
INPUT_DIRS = [
    r"C:/ai_x/source/Project2nd/금융법령수집/금융문제수집/금융법령",
    r"C:/ai_x/source/Project2nd/금융법령수집/금융문제수집/기업법령",
    r"C:/ai_x/source/Project2nd/금융법령수집/금융문제수집/보험법령",
    r"C:/ai_x/source/Project2nd/금융법령수집/금융문제수집/상법투자자산증권주식법령",
    r"C:/ai_x/source/Project2nd/금융법령수집/금융문제수집/은행법령",
    r"C:/ai_x/source/Project2nd/금융법령수집/금융문제수집/조합법령",
]

INDEX_NAME = "codedoc-law-upstage"   # 요청하신 인덱스 이름
EMBED_MODEL = os.getenv("UPSTAGE_EMBEDDING_MODEL", "solar-embedding-1-large")
DIMENSION = 4096                     # Upstage solar-embedding-1-large 차원
METRIC = "cosine"
NAMESPACE = os.getenv("NAMESPACE", "")  # 필요시 설정

# 청크 사이즈/오버랩 (문자 단위)
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 200

# 업서트 배치 크기
BATCH_SIZE = 64

# Pinecone 서버리스 스펙 (환경변수 없으면 기본값)
PC_CLOUD = os.getenv("PINECONE_CLOUD", "aws")
PC_REGION = os.getenv("PINECONE_REGION", "us-east-1")

# ===================== 유틸 함수들 =====================

def norm_ws(s: str) -> str:
    return " ".join((s or "").split())

def read_docx(path: str) -> Tuple[str, str]:
    """
    .docx 본문을 읽어서 (title, full_text) 반환
    제목은 문서 첫 Heading 또는 첫 문장 정도로 추출
    """
    doc = Document(path)
    paras = [norm_ws(p.text) for p in doc.paragraphs if norm_ws(p.text)]
    title = ""
    for p in doc.paragraphs:
        if p.style and "Heading" in str(p.style.name):
            title = norm_ws(p.text)
            break
    if not title:
        # 첫 문장 일부를 타이틀 대용으로
        title = (paras[0][:60] if paras else os.path.basename(path))
    full = "\n".join(paras)
    return title, full

def chunk_text(text: str, max_len: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    단순 문자 기반 청크 (한국어 규정형 텍스트에 무해)
    """
    text = norm_ws(text)
    n = len(text)
    chunks: List[str] = []
    if n == 0:
        return chunks
    start = 0
    while start < n:
        end = min(n, start + max_len)
        chunk = text[start:end]
        # 문장 경계 배려: 가능한 마지막 마침표까지
        if end < n:
            cut = chunk.rfind("다.")
            if cut != -1 and (end - (start + cut) < 80):
                chunk = chunk[:cut + 2]
                end = start + len(chunk)
        chunks.append(chunk)
        if end >= n:
            break
        start = max(0, end - overlap)
    return chunks

def id_for(path: str, chunk_idx: int) -> str:
    h = hashlib.md5(path.encode("utf-8")).hexdigest()
    return f"{h}-{chunk_idx:05d}"

def iter_docx_files(dirs: List[str]) -> Iterator[str]:
    for d in dirs:
        for p in glob.glob(os.path.join(d, "**", "*.docx"), recursive=True):
            yield p

def backoff_sleep(attempt: int):
    # 0, 1, 2, ... -> 0.5, 1.0, 2.0, 4.0, ... (최대 8초)
    sec = min(8.0, 0.5 * (2 ** max(0, attempt)))
    time.sleep(sec)

# ===================== 메인 파이프라인 =====================

def ensure_index(pc: Pinecone, name: str, dimension: int, metric: str):
    # 인덱스가 없으면 생성
    names = [x["name"] if isinstance(x, dict) else getattr(x, "name", None) for x in pc.list_indexes()]
    if name not in names:
        print(f"[Pinecone] creating index: {name} (dim={dimension}, metric={metric}, {PC_CLOUD}/{PC_REGION})")
        pc.create_index(
            name=name,
            dimension=dimension,
            metric=metric,
            spec=ServerlessSpec(cloud=PC_CLOUD, region=PC_REGION)
        )
    else:
        # 차원 불일치 시에는 수동으로 삭제 후 재생성 필요
        print(f"[Pinecone] index exists: {name}")

def embed_batches(emb: UpstageEmbeddings, texts: List[str]) -> List[List[float]]:
    # Upstage LangChain 래퍼는 내부적으로 배치 처리 지원
    return emb.embed_documents(texts)

def main():
    # 키 체크
    up_key = os.getenv("UPSTAGE_API_KEY")
    pc_key = os.getenv("PINECONE_API_KEY")
    if not up_key:
        print("ERROR: UPSTAGE_API_KEY 가 필요합니다.", file=sys.stderr)
        sys.exit(1)
    if not pc_key:
        print("ERROR: PINECONE_API_KEY 가 필요합니다.", file=sys.stderr)
        sys.exit(1)

    # Upstage 임베딩
    emb = UpstageEmbeddings(model=EMBED_MODEL)
    print(f"[Upstage] model={EMBED_MODEL}")

    # Pinecone 클라이언트 + 인덱스 준비
    pc = Pinecone(api_key=pc_key)
    ensure_index(pc, INDEX_NAME, DIMENSION, METRIC)
    index = pc.Index(INDEX_NAME)

    # 파일 수집
    files = list(iter_docx_files(INPUT_DIRS))
    if not files:
        print("[INFO] .docx 파일을 찾지 못했습니다.")
        return
    print(f"[INFO] .docx files: {len(files)}")

    total_chunks = 0
    upserted = 0

    batch_ids: List[str] = []
    batch_vecs: List[List[float]] = []
    batch_metas: List[Dict] = []

    def flush():
        nonlocal upserted, batch_ids, batch_vecs, batch_metas
        if not batch_ids:
            return
        # Pinecone 업서트 (리트라이 백오프)
        attempt = 0
        while True:
            try:
                vectors = []
                for i in range(len(batch_ids)):
                    vectors.append({
                        "id": batch_ids[i],
                        "values": batch_vecs[i],
                        "metadata": batch_metas[i],
                    })
                index.upsert(vectors=vectors, namespace=NAMESPACE)
                upserted += len(batch_ids)
                print(f"[UPSERT] +{len(batch_ids)} (total {upserted})")
                batch_ids.clear()
                batch_vecs.clear()
                batch_metas.clear()
                break
            except Exception as e:
                print(f"[WARN] upsert failed (attempt {attempt}): {e}")
                attempt += 1
                backoff_sleep(attempt)

    for fi, path in enumerate(files, 1):
        try:
            title, text = read_docx(path)
            if not text.strip():
                continue
            chunks = chunk_text(text, CHUNK_SIZE, CHUNK_OVERLAP)
            if not chunks:
                continue
            total_chunks += len(chunks)

            # 임베딩 배치
            for i in range(0, len(chunks), BATCH_SIZE):
                batch = chunks[i : i + BATCH_SIZE]
                vecs = embed_batches(emb, batch)

                for j, v in enumerate(vecs):
                    cid = id_for(path, i + j)
                    meta = {
                        "source_path": path,
                        "chunk_id": i + j,
                        "title": title,
                        "doc_type": "docx",
                        # Pinecone 메타 용량 보호: 텍스트는 일부만
                        "text_preview": batch[j][:200],
                    }
                    batch_ids.append(cid)
                    batch_vecs.append(v)
                    batch_metas.append(meta)

                if len(batch_ids) >= BATCH_SIZE:
                    flush()

            # 진행 로그
            if fi % 20 == 0:
                print(f"[PROGRESS] files {fi}/{len(files)} | chunks so far {total_chunks} | upserted {upserted}")

        except Exception as e:
            print(f"[ERROR] {path}: {e}")

    flush()
    print(f"[DONE] files={len(files)} chunks={total_chunks} upserted={upserted} index={INDEX_NAME} ns='{NAMESPACE}'")


if __name__ == "__main__":
    main()
